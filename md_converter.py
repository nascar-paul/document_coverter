import markdown
from docx import Document
from bs4 import BeautifulSoup
import os

# Function to convert Markdown text to Word document
def markdown_to_word(input_file, output_file):
    # Read the markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_text = f.read()

    # Extract the filename without extension for the Word document title
    filename = os.path.basename(input_file)
    file_title = os.path.splitext(filename)[0]

    # Convert markdown text to HTML
    html = markdown.markdown(markdown_text)

    # Parse the HTML
    soup = BeautifulSoup(html, 'html.parser')

    # Create a new Word document
    doc = Document()

    # Add the filename as the title of the Word document
    doc.add_heading(f"Converted from: {file_title}", level=1)

    # Go through the HTML tags and add them to the Word document
    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text(), style='ListBullet')

    # Save the document
    doc.save(output_file)

# Ask user for the markdown file name and output file name
input_file = input("Enter the name of the Markdown file (with extension, e.g., 'example_markdown.md'): ")
output_file = input("Enter the desired name of the output Word file (with extension, e.g., 'output.docx'): ")

# Convert markdown to Word document
markdown_to_word(input_file, output_file)
print(f"Markdown from {input_file} has been converted to {output_file}")
